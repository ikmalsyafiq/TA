import base64
import hashlib
import io
import os
import re
from pathlib import Path
from datetime import datetime

import streamlit as st
from PIL import Image
import requests
try:
    import truststore
except ModuleNotFoundError:
    truststore = None

from openai import OpenAI
from openai import APIConnectionError, APIStatusError, AuthenticationError, BadRequestError, RateLimitError


st.set_page_config(page_title="Commodities TA Analyzer", layout="wide")

if truststore is not None:
    truststore.inject_into_ssl()

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def to_data_url(uploaded_file, max_dimension: int = 1800, jpeg_quality: int = 85) -> str:
    uploaded_file.seek(0)
    image = Image.open(uploaded_file)
    image = image.convert("RGB")

    width, height = image.size
    longest_side = max(width, height)
    if longest_side > max_dimension:
        scale = max_dimension / longest_side
        resized = (int(width * scale), int(height * scale))
        image = image.resize(resized, Image.Resampling.LANCZOS)

    buffer = io.BytesIO()
    image.save(buffer, format="JPEG", quality=jpeg_quality, optimize=True)
    encoded = base64.b64encode(buffer.getvalue()).decode("utf-8")
    return f"data:image/jpeg;base64,{encoded}"


def build_prompt(
    instrument: str,
    manual_support: str = "",
    manual_resistance: str = "",
    additional_input: str = "",
    uploaded_file_names: list[str] | None = None,
) -> str:
    instrument_line = instrument.strip() or "Instrument not specified"
    support_line = manual_support.strip()
    resistance_line = manual_resistance.strip()
    additional_line = additional_input.strip()
    uploaded_file_names = uploaded_file_names or []
    files_line = ", ".join(uploaded_file_names) if uploaded_file_names else "Not provided"
    level_guidance = ""
    if support_line or resistance_line:
        level_guidance = f"""
User-provided key levels (treat as priority anchors unless clearly invalidated by current chart action):
- Support: {support_line or 'Not provided'}
- Resistance: {resistance_line or 'Not provided'}
""".strip()

    additional_guidance = ""
    if additional_line:
        additional_guidance = f"""
Additional user context/instructions (optional):
{additional_line}
""".strip()

    return f"""
You are a senior technical analyst specializing in energy and commodity markets.
Analyze the uploaded chart screenshots and produce a professional trading-desk technical briefing.

Instrument: {instrument_line}
Date context: {datetime.now().strftime('%Y-%m-%d')}
Uploaded screenshot file names: {files_line}

{level_guidance}

{additional_guidance}

Required structure:
1) Instrument + Contract
2) Timeframe map + Bias by detected timeframe
    - First detect available timeframes from chart labels and file names.
    - Then report bias in sequence from lowest to highest timeframe detected.
    - Do NOT force missing frames (e.g., if only 1H/2H/4H/D are present, use those only).
3) Key Levels: Support (3) and Resistance (3)
4) Primary Trade Idea (trend-aligned): direction, entry, target, stretch target, stop, reason, invalidation
5) Structure by detected timeframe:
    - For each detected timeframe, provide 2-4 concise observations + one takeaway.
6) Momentum & Indicators by detected timeframe: RSI, Stochastic, MACD, and ATR/ADX (if visible) + takeaway
9) Bottom Line
10) Alternative Trade Setup (lower probability)

Rules:
- Prioritize price action and structure over indicators.
- Use concrete levels from charts where visible.
- Write concise, professional, actionable language.
- If timeframe detection is uncertain for any image, make the best estimate and proceed.
- If ATR or ADX is visible, interpret:
    - ATR as volatility/expansion vs contraction context
    - ADX as trend strength (e.g., weak, developing, strong)
""".strip()


def generate_analysis(client: OpenAI, model: str, prompt: str, image_data_urls: list[str]) -> str:
    content = [{"type": "text", "text": prompt}]
    for url in image_data_urls:
        content.append({"type": "image_url", "image_url": {"url": url}})

    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": content}],
        temperature=0.2,
        max_tokens=2000,
        timeout=120,
    )
    if not response.choices:
        return "No analysis generated."
    message_content = response.choices[0].message.content
    if isinstance(message_content, str):
        return message_content
    if isinstance(message_content, list):
        text_parts = [
            item.get("text", "")
            for item in message_content
            if isinstance(item, dict) and item.get("type") == "text"
        ]
        return "\n".join(part for part in text_parts if part).strip() or "No analysis generated."
    return "No analysis generated."


def build_client(api_key: str, base_url: str = "") -> OpenAI:
    if base_url:
        return OpenAI(api_key=api_key, base_url=base_url)
    return OpenAI(api_key=api_key)


def normalize_base_url(url: str) -> str:
    return (url or "").strip().rstrip("/")


def github_base_url_candidates(base_url: str) -> list[str]:
    candidates: list[str] = []
    primary = normalize_base_url(base_url)
    if primary:
        candidates.append(primary)

    known = [
        "https://models.github.ai/inference",
        "https://models.inference.ai.azure.com",
        "https://models.github.ai",
    ]
    for item in known:
        if item not in candidates:
            candidates.append(item)
    return candidates


def github_model_candidates(model: str) -> list[str]:
    candidates: list[str] = []
    primary = (model or "").strip()
    if primary:
        candidates.append(primary)

    known = [
        "auto",
        "openai/gpt-4o",
        "gpt-4o",
        "openai/gpt-4.1",
        "gpt-4.1",
        "openai/gpt-4o-mini",
    ]
    for item in known:
        if item not in candidates:
            candidates.append(item)
    return candidates


def is_unknown_model_error(error: APIStatusError) -> bool:
    message = str(error).lower()
    return error.status_code == 400 and ("unknown_model" in message or "unknown model" in message)


def test_provider_connection(client: OpenAI, model: str) -> None:
    client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": "ping"}],
        max_tokens=16,
        timeout=30,
    )


def test_https_reachability(url: str) -> tuple[bool, str]:
    try:
        response = requests.get(url, timeout=15)
        return True, f"Reachable. HTTP {response.status_code}"
    except Exception as error:
        return False, str(error)


def get_working_client(provider: str, api_key: str, base_url: str, model: str) -> tuple[OpenAI, str, str]:
    if provider != "GitHub Models (Copilot-style)":
        return build_client(api_key, ""), "", model

    last_error = None
    for candidate in github_base_url_candidates(base_url):
        client = build_client(api_key, candidate)
        for candidate_model in github_model_candidates(model):
            try:
                test_provider_connection(client, candidate_model)
                return client, candidate, candidate_model
            except APIStatusError as error:
                last_error = error
                if error.status_code == 404:
                    break
                if is_unknown_model_error(error):
                    continue
                raise
            except Exception as error:
                last_error = error
                continue

    if last_error:
        raise last_error
    raise RuntimeError("Unable to find a working GitHub Models endpoint/model combination.")


def persist_uploads(uploaded_files) -> tuple[list[str], list[Path], str]:
    data_urls: list[str] = []
    saved_paths: list[Path] = []
    hasher = hashlib.sha256()

    batch_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    for index, uploaded in enumerate(uploaded_files, start=1):
        uploaded.seek(0)
        raw_bytes = uploaded.read()
        uploaded.seek(0)

        hasher.update(raw_bytes)
        hasher.update(uploaded.name.encode("utf-8"))

        suffix = Path(uploaded.name).suffix or ".png"
        safe_name = f"{batch_id}_{index}{suffix}"
        save_path = UPLOAD_DIR / safe_name
        save_path.write_bytes(raw_bytes)
        saved_paths.append(save_path)

        data_urls.append(to_data_url(uploaded))

    return data_urls, saved_paths, hasher.hexdigest()


def analysis_to_docx_bytes(title: str, analysis: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as error:
        raise RuntimeError(
            "Word export requires python-docx. Install dependency 'python-docx'."
        ) from error

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    doc.add_heading(title, level=1)

    def normalize_text(text: str) -> str:
        cleaned = text.strip()
        cleaned = cleaned.replace("\t", " ")
        cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned)
        cleaned = re.sub(r"__(.*?)__", r"\1", cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned)
        cleaned = cleaned.replace("\u2014", "-")
        return cleaned.strip()

    blank_streak = 0
    for line in analysis.splitlines():
        stripped = normalize_text(line)
        if not stripped:
            blank_streak += 1
            if blank_streak <= 1:
                doc.add_paragraph("")
            continue
        blank_streak = 0

        section_match = re.match(r"^(?:#{1,6}\s*)?(\d+)\)\s+(.+)$", stripped)
        if section_match:
            doc.add_heading(f"{section_match.group(1)}) {section_match.group(2).strip()}", level=2)
        elif re.match(r"^#{1,6}\s+", stripped):
            heading_text = re.sub(r"^#{1,6}\s+", "", stripped).strip()
            doc.add_heading(heading_text, level=2)
        elif stripped.endswith(":") and len(stripped) <= 80:
            doc.add_heading(stripped[:-1], level=3)
        elif stripped.startswith("-") or stripped.startswith("•"):
            bullet_text = stripped.lstrip("-• ").strip()
            if bullet_text:
                doc.add_paragraph(bullet_text, style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            numbered_text = re.sub(r"^\d+\.\s+", "", stripped)
            doc.add_paragraph(numbered_text.strip(), style="List Number")
        else:
            doc.add_paragraph(stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


st.title("Commodities TA Analyzer")
st.caption("Upload chart screenshots (1H/1D/1W) and generate a structured technical analysis report.")

st.info("Upload screenshots and the app auto-runs analysis. Screenshots are saved locally in the uploads folder.")

provider = st.selectbox("Provider", ["GitHub Models (Copilot-style)", "OpenAI"], index=0)

if provider == "GitHub Models (Copilot-style)":
    api_key = st.text_input(
        "GitHub Token",
        type="password",
        value=os.getenv("GITHUB_TOKEN", ""),
        help="Use a GitHub token with access to GitHub Models.",
    )
    base_url = st.text_input(
        "GitHub Models Base URL",
        value=os.getenv("GITHUB_MODELS_BASE_URL", "https://models.github.ai/inference"),
        help="Use https://models.inference.ai.azure.com if your network blocks models.github.ai.",
    )
    default_model = os.getenv("GITHUB_MODEL", "auto")
    https_test_url = "https://models.github.ai/inference"
    connection_label = "Test GitHub Models Connection"
else:
    api_key = st.text_input(
        "OpenAI API Key",
        type="password",
        value=os.getenv("OPENAI_API_KEY", ""),
        help="Required for analysis generation. Used only during this session.",
    )
    base_url = ""
    default_model = os.getenv("OPENAI_MODEL", "gpt-4.1")
    https_test_url = "https://api.openai.com/v1/models"
    connection_label = "Test OpenAI Connection"

with st.expander("Advanced (optional)"):
    model_override = st.text_input(
        "Model override",
        value="",
        help="Leave empty to use automatic/default model selection.",
    )

model = model_override.strip() or default_model

if st.button(connection_label):
    if not api_key:
        st.error("Please provide an API key/token first.")
    else:
        try:
            with st.spinner("Testing connection..."):
                if provider == "GitHub Models (Copilot-style)":
                    client, resolved_base_url, resolved_model = get_working_client(provider, api_key, base_url, model)
                    if model_override.strip() and model.strip() != resolved_model:
                        st.warning(f"Model '{model}' unavailable. Using fallback model: {resolved_model}")
                else:
                    client = build_client(api_key, "")
                    test_provider_connection(client, model)
            st.success("Connection test passed.")
        except AuthenticationError as error:
            st.error(f"Authentication failed: {error}")
        except APIConnectionError as error:
            st.error(f"Connection error: {error}")
            st.info("Verify outbound HTTPS access and corporate proxy settings for the selected provider endpoint.")
        except APIStatusError as error:
            st.error(f"API status error ({error.status_code}): {error}")
        except Exception as error:
            st.error(f"Connection test failed: {error}")

if st.button("Test HTTPS (SSL/Proxy)"):
    ok, message = test_https_reachability(https_test_url)
    if ok:
        st.success(f"HTTPS check passed: {message}")
    else:
        st.error(f"HTTPS check failed: {message}")
        st.info("If you are on a corporate network, verify proxy settings and outbound SSL inspection rules for this endpoint.")

instrument = st.text_input("Instrument + Contract", value="Dutch TTF Natural Gas Futures (ICE Endex)")

st.subheader("Manual Key Levels (Optional)")
manual_support = st.text_input(
    "Support levels",
    value="",
    help="Optional. Example: 113.0 / 109.5 / 99.0",
)
manual_resistance = st.text_input(
    "Resistance levels",
    value="",
    help="Optional. Example: 119.5 / 125.0 / 139.0",
)

additional_input = st.text_area(
    "Additional Analyst Input (Optional)",
    value="",
    help="Optional extra instructions/context for the agent, e.g. event risk, preferred trade style, or key assumptions.",
)

uploads = st.file_uploader(
    "Upload screenshots (PNG/JPG)",
    type=["png", "jpg", "jpeg"],
    accept_multiple_files=True,
)

if uploads:
    cols = st.columns(min(3, len(uploads)))
    for idx, up in enumerate(uploads):
        cols[idx % len(cols)].image(up, caption=up.name, width="stretch")

if "last_upload_signature" not in st.session_state:
    st.session_state.last_upload_signature = ""
if "latest_analysis" not in st.session_state:
    st.session_state.latest_analysis = ""
if "latest_doc_name" not in st.session_state:
    st.session_state.latest_doc_name = ""
if "latest_doc_bytes" not in st.session_state:
    st.session_state.latest_doc_bytes = b""
if "latest_saved_paths" not in st.session_state:
    st.session_state.latest_saved_paths = []
if "active_provider" not in st.session_state:
    st.session_state.active_provider = ""
if "active_endpoint" not in st.session_state:
    st.session_state.active_endpoint = ""
if "active_model" not in st.session_state:
    st.session_state.active_model = ""

if uploads and api_key:
    try:
        data_urls, saved_paths, upload_signature = persist_uploads(uploads)
        st.session_state.latest_saved_paths = [str(path) for path in saved_paths]

        if upload_signature != st.session_state.last_upload_signature:
            with st.spinner("New screenshots detected. Running analysis automatically..."):
                if provider == "GitHub Models (Copilot-style)":
                    client, resolved_base_url, resolved_model = get_working_client(provider, api_key, base_url, model)
                else:
                    client = build_client(api_key, "")
                    resolved_model = model
                    resolved_base_url = "https://api.openai.com"
                uploaded_file_names = [upload.name for upload in uploads]
                prompt = build_prompt(
                    instrument,
                    manual_support,
                    manual_resistance,
                    additional_input,
                    uploaded_file_names,
                )
                analysis = generate_analysis(client, resolved_model, prompt, data_urls)
                doc_name = f"TA_Analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                try:
                    doc_bytes = analysis_to_docx_bytes(instrument, analysis)
                except Exception:
                    doc_bytes = b""

            st.session_state.last_upload_signature = upload_signature
            st.session_state.latest_analysis = analysis
            st.session_state.latest_doc_name = doc_name
            st.session_state.latest_doc_bytes = doc_bytes
            st.session_state.active_provider = provider
            st.session_state.active_endpoint = resolved_base_url
            st.session_state.active_model = resolved_model
    except AuthenticationError as error:
        st.error(f"Authentication failed: {error}")
    except RateLimitError as error:
        st.error(f"Rate limit reached: {error}")
        st.info("Retry in a moment, reduce request frequency, or switch to a different model.")
    except APIConnectionError as error:
        st.error(f"Connection error: {error}")
        st.info("Verify outbound HTTPS access and corporate proxy settings for the selected provider endpoint.")
    except BadRequestError as error:
        st.error(f"Request rejected: {error}")
        st.info("Try fewer screenshots or lower-resolution images. The app auto-compresses images, but very large batches can still fail.")
    except APIStatusError as error:
        st.error(f"API status error ({error.status_code}): {error}")
        if error.status_code == 404 and provider == "GitHub Models (Copilot-style)":
            st.info("404 usually means endpoint mismatch. Try base URL https://models.inference.ai.azure.com or run the connection test to auto-fallback.")
    except Exception as error:
        st.error(f"Failed to generate analysis: {error}")

if uploads and not api_key:
    if provider == "GitHub Models (Copilot-style)":
        st.warning("Please provide a GitHub token to run automatic analysis.")
    else:
        st.warning("Please provide an OpenAI API key to run automatic analysis.")

if st.session_state.latest_saved_paths:
    st.subheader("Saved Screenshots")
    for path in st.session_state.latest_saved_paths:
        st.write(path)

if st.session_state.latest_analysis:
    st.subheader("Active Runtime")
    st.write(f"Provider: {st.session_state.active_provider}")
    st.write(f"Endpoint: {st.session_state.active_endpoint}")
    st.write(f"Model: {st.session_state.active_model}")

    st.subheader("Analysis")
    st.write(st.session_state.latest_analysis)
    if st.session_state.latest_doc_bytes:
        st.download_button(
            "Download Word Report",
            data=st.session_state.latest_doc_bytes,
            file_name=st.session_state.latest_doc_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.warning("Word export is unavailable because python-docx is not installed in this environment.")

st.divider()
st.markdown("""
**How to use**
- Select provider.
- Paste provider token/key.
- Upload 1H, 1D, and 1W screenshots.
- Analysis runs automatically and appears below.
""")
